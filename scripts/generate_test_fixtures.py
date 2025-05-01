import os

def ensure_dir(path):
    if not os.path.exists(path):
        os.makedirs(path)

def create_sample_pdf(filepath):
    # Uses reportlab to create a minimal PDF
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.pdfgen import canvas
    except ImportError:
        raise RuntimeError("reportlab is required: pip install reportlab")

    c = canvas.Canvas(filepath, pagesize=letter)
    c.drawString(100, 750, "Introduction to Science - PDF fixture for AeroLearn AI integration test.")
    c.save()

def create_sample_docx(filepath):
    # Uses python-docx to create a minimal DOCX with a table
    try:
        import docx
    except ImportError:
        raise RuntimeError("python-docx is required: pip install python-docx")

    doc = docx.Document()
    doc.add_heading("Sample Report", 0)
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = 'Header 1'
    table.cell(0, 1).text = 'Header 2'
    table.cell(1, 0).text = 'Cell 1'
    table.cell(1, 1).text = 'Cell 2'
    doc.save(filepath)

if __name__ == "__main__":
    # Destination directory as in test_config
    TEST_DIR = "tests/fixtures/sample_content/"
    ensure_dir(TEST_DIR)
    pdf_path = os.path.join(TEST_DIR, "sample_lecture_notes.pdf")
    docx_path = os.path.join(TEST_DIR, "sample_report.docx")

    create_sample_pdf(pdf_path)
    create_sample_docx(docx_path)
    print(f"Fixtures created: {pdf_path}, {docx_path}")